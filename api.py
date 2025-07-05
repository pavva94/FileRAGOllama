from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Optional
import tempfile
import os
from pathlib import Path
import asyncio
import uuid
from datetime import datetime
import json
import logging

# Database imports
import asyncpg
from sqlalchemy import create_engine, Column, String, DateTime, Integer, Text, Float
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session
from sqlalchemy.dialects.postgresql import UUID
import databases

# Azure imports
from azure.storage.blob import BlobServiceClient
from azure.core.exceptions import AzureError

# Gemini imports
import google.generativeai as genai
from sentence_transformers import SentenceTransformer

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration
DATABASE_URL = os.getenv("DATABASE_URL", "postgresql://rag_user:rag_password@postgres:5432/rag_db")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
AZURE_STORAGE_CONNECTION_STRING = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
AZURE_STORAGE_CONTAINER_NAME = os.getenv("AZURE_STORAGE_CONTAINER_NAME", "rag-documents")

# Initialize services
if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)
    gemini_model = genai.GenerativeModel('gemini-pro')
else:
    gemini_model = None
    logger.warning("GOOGLE_API_KEY not set. Gemini functionality will be disabled.")

if AZURE_STORAGE_CONNECTION_STRING:
    try:
        blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        # Create container if it doesn't exist
        try:
            blob_service_client.create_container(AZURE_STORAGE_CONTAINER_NAME)
        except:
            pass  # Container might already exist
    except Exception as e:
        logger.error(f"Failed to initialize Azure Blob Storage: {e}")
        blob_service_client = None
else:
    blob_service_client = None
    logger.warning("AZURE_STORAGE_CONNECTION_STRING not set. Using local storage.")

# Initialize embedding model
try:
    embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
except Exception as e:
    logger.error(f"Failed to load embedding model: {e}")
    embedding_model = None

# Database setup
database = databases.Database(DATABASE_URL)
engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

Base = declarative_base()


class Document(Base):
    __tablename__ = "documents"

    id = Column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    filename = Column(String, nullable=False)
    file_size = Column(Integer, nullable=False)
    upload_date = Column(DateTime, default=datetime.utcnow)
    chunk_count = Column(Integer, default=0)
    azure_blob_url = Column(String, nullable=True)


class DocumentChunk(Base):
    __tablename__ = "document_chunks"

    id = Column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    document_id = Column(UUID(as_uuid=True), nullable=False)
    chunk_index = Column(Integer, nullable=False)
    content = Column(Text, nullable=False)
    embedding = Column(Text, nullable=True)  # JSON string of embedding vector


# Create tables
Base.metadata.create_all(bind=engine)

app = FastAPI(title="Simple RAG API with Azure & Gemini", version="2.0.0")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# Dependency to get DB session
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# Pydantic models
class QuestionRequest(BaseModel):
    question: str
    max_results: int = 5


class AnswerResponse(BaseModel):
    answer: str
    sources: List[str]
    confidence: float


class FileInfo(BaseModel):
    id: str
    filename: str
    file_size: int
    upload_date: str
    chunk_count: int


class UploadResponse(BaseModel):
    file_id: str
    filename: str
    chunk_count: int
    file_size: int


# Utility functions
def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks"""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start = end - overlap
        if start >= len(text):
            break
    return chunks


def extract_text_from_file(file_path: str) -> str:
    """Extract text from various file formats"""
    try:
        file_extension = Path(file_path).suffix.lower()

        if file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_extension == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text()
                    return text
            except ImportError:
                raise HTTPException(status_code=500, detail="PyPDF2 not installed for PDF processing")
        elif file_extension == '.docx':
            try:
                from docx import Document
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except ImportError:
                raise HTTPException(status_code=500, detail="python-docx not installed for DOCX processing")
        elif file_extension == '.md':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_extension}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting text: {str(e)}")


async def upload_to_azure(file_content: bytes, filename: str) -> str:
    """Upload file to Azure Blob Storage"""
    if not blob_service_client:
        return None

    try:
        blob_name = f"{uuid.uuid4()}_{filename}"
        blob_client = blob_service_client.get_blob_client(
            container=AZURE_STORAGE_CONTAINER_NAME,
            blob=blob_name
        )

        await asyncio.to_thread(blob_client.upload_blob, file_content, overwrite=True)
        return blob_client.url
    except Exception as e:
        logger.error(f"Failed to upload to Azure: {e}")
        return None


def get_embedding(text: str) -> List[float]:
    """Generate embedding for text"""
    if not embedding_model:
        return []

    try:
        embedding = embedding_model.encode(text)
        return embedding.tolist()
    except Exception as e:
        logger.error(f"Failed to generate embedding: {e}")
        return []


def calculate_similarity(embedding1: List[float], embedding2: List[float]) -> float:
    """Calculate cosine similarity between two embeddings"""
    if not embedding1 or not embedding2:
        return 0.0

    try:
        import numpy as np
        from scipy.spatial.distance import cosine
        return 1 - cosine(embedding1, embedding2)
    except ImportError:
        return 0.0


async def generate_with_gemini(prompt: str) -> str:
    """Generate response using Gemini"""
    if not gemini_model:
        return None

    try:
        response = await asyncio.to_thread(gemini_model.generate_content, prompt)
        return response.text
    except Exception as e:
        logger.error(f"Gemini generation failed: {e}")
        return None


# API Endpoints
@app.on_event("startup")
async def startup():
    await database.connect()


@app.on_event("shutdown")
async def shutdown():
    await database.disconnect()


@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "database": "connected" if database.is_connected else "disconnected",
        "gemini": "available" if gemini_model else "unavailable",
        "azure_storage": "available" if blob_service_client else "unavailable",
        "embeddings": "available" if embedding_model else "unavailable"
    }


@app.post("/upload", response_model=UploadResponse)
async def upload_file(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """Upload a file to the RAG system"""

    # Check file type
    allowed_extensions = ['.txt', '.pdf', '.md', '.docx']
    file_extension = Path(file.filename).suffix.lower()

    if file_extension not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
        )

    try:
        # Read file content
        content = await file.read()

        # Save to temporary file for processing
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
            tmp_file.write(content)
            tmp_file_path = tmp_file.name

        # Extract text
        text_content = extract_text_from_file(tmp_file_path)

        # Upload to Azure (optional)
        azure_url = await upload_to_azure(content, file.filename)

        # Create document record
        doc_id = uuid.uuid4()
        document = Document(
            id=doc_id,
            filename=file.filename,
            file_size=len(content),
            azure_blob_url=azure_url
        )

        # Create chunks
        chunks = chunk_text(text_content)
        chunk_objects = []

        for i, chunk in enumerate(chunks):
            embedding = get_embedding(chunk)
            chunk_obj = DocumentChunk(
                document_id=doc_id,
                chunk_index=i,
                content=chunk,
                embedding=json.dumps(embedding) if embedding else None
            )
            chunk_objects.append(chunk_obj)

        # Save to database
        db.add(document)
        db.add_all(chunk_objects)
        document.chunk_count = len(chunks)
        db.commit()

        # Clean up temp file
        os.unlink(tmp_file_path)

        return UploadResponse(
            file_id=str(doc_id),
            filename=file.filename,
            chunk_count=len(chunks),
            file_size=len(content)
        )

    except Exception as e:
        db.rollback()
        if 'tmp_file_path' in locals():
            try:
                os.unlink(tmp_file_path)
            except:
                pass
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/files", response_model=List[FileInfo])
async def get_files(db: Session = Depends(get_db)):
    """Get list of uploaded files"""
    try:
        documents = db.query(Document).all()
        return [
            FileInfo(
                id=str(doc.id),
                filename=doc.filename,
                file_size=doc.file_size,
                upload_date=doc.upload_date.isoformat(),
                chunk_count=doc.chunk_count
            )
            for doc in documents
        ]
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/files/{file_id}")
async def delete_file(file_id: str, db: Session = Depends(get_db)):
    """Delete a file"""
    try:
        # Delete chunks first
        db.query(DocumentChunk).filter(DocumentChunk.document_id == file_id).delete()

        # Delete document
        doc = db.query(Document).filter(Document.id == file_id).first()
        if not doc:
            raise HTTPException(status_code=404, detail="File not found")

        db.delete(doc)
        db.commit()

        return {"message": "File deleted successfully"}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/ask", response_model=AnswerResponse)
async def ask_question(request: QuestionRequest, db: Session = Depends(get_db)):
    """Ask a question to the RAG system"""
    try:
        # Get question embedding
        question_embedding = get_embedding(request.question)

        # Find similar chunks
        chunks = db.query(DocumentChunk).all()
        similar_chunks = []

        for chunk in chunks:
            if chunk.embedding:
                chunk_embedding = json.loads(chunk.embedding)
                similarity = calculate_similarity(question_embedding, chunk_embedding)
                similar_chunks.append((chunk, similarity))

        # Sort by similarity and take top results
        similar_chunks.sort(key=lambda x: x[1], reverse=True)
        top_chunks = similar_chunks[:request.max_results]

        if not top_chunks:
            return AnswerResponse(
                answer="I don't have enough information to answer your question.",
                sources=[],
                confidence=0.0
            )

        # Prepare context
        context = "\n\n".join([chunk.content for chunk, _ in top_chunks])

        # Get sources (document filenames)
        doc_ids = [chunk.document_id for chunk, _ in top_chunks]
        documents = db.query(Document).filter(Document.id.in_(doc_ids)).all()
        sources = list(set([doc.filename for doc in documents]))

        # Calculate confidence (average similarity)
        confidence = sum([sim for _, sim in top_chunks]) / len(top_chunks)

        # Generate answer with Gemini
        answer = None
        if gemini_model:
            prompt = f"""Based on the following context, please answer the question clearly and concisely.

Context:
{context}

Question: {request.question}

Please provide a helpful answer based on the context. If the context doesn't contain enough information, please say so.

Answer:"""

            answer = await generate_with_gemini(prompt)

        # Fallback to basic answer if Gemini fails
        if not answer:
            answer = f"Based on the available documents, here's what I found:\n\n{context[:500]}..."

        return AnswerResponse(
            answer=answer,
            sources=sources,
            confidence=confidence
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/gemini/status")
async def gemini_status():
    """Check Gemini status"""
    return {
        "available": gemini_model is not None,
        "model": "gemini-pro" if gemini_model else None
    }


@app.get("/")
async def root():
    """Root endpoint"""
    return {"message": "Simple RAG API with Azure & Gemini is running!"}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=3000)